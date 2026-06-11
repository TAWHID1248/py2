"""python-docx Word template filling + PDF conversion."""
from pathlib import Path
from app.models.contact import Contact
from app.services.content.personalization import build_context
from app.core.logger import get_logger

log = get_logger(__name__)


def fill_template(template_path: Path, contact: Contact, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt
    import copy

    doc = Document(str(template_path))
    ctx = build_context(contact)

    def _replace_in_run(run):
        for key, val in ctx.items():
            placeholder = f"{{{key}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, val or "")

    def _replace_paragraph(para):
        # Handle split placeholders across runs
        full_text = "".join(r.text for r in para.runs)
        for key, val in ctx.items():
            full_text = full_text.replace(f"{{{key}}}", val or "")
        if para.runs:
            para.runs[0].text = full_text
            for r in para.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        _replace_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_paragraph(para)

    # Replace in headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_paragraph(para)
        for para in section.footer.paragraphs:
            _replace_paragraph(para)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.debug("Word doc: %s", output_path.name)
    return output_path


def docx_to_pdf(docx_path: Path, output_path: Path) -> Path:
    """Convert .docx → .pdf using docx2pdf (requires MS Word or LibreOffice)."""
    try:
        import docx2pdf
        docx2pdf.convert(str(docx_path), str(output_path))
        return output_path
    except Exception as e:
        log.warning("docx2pdf failed (%s) — falling back to HTML→PDF", e)
        # Fallback: read doc text → HTML → PDF
        from docx import Document
        from app.services.document.pdf_service import html_to_pdf

        doc = Document(str(docx_path))
        lines = [f"<p>{para.text}</p>" for para in doc.paragraphs if para.text.strip()]
        html = "<html><body>" + "".join(lines) + "</body></html>"
        return html_to_pdf(html, output_path)


def create_from_html(html: str, output_path: Path) -> Path:
    """Create a .docx from raw HTML using python-docx (basic conversion)."""
    import re
    from docx import Document

    doc = Document()
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").strip()

    for line in text.split("\n"):
        doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
